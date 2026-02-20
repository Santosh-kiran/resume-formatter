import streamlit as st
import fitz  # PyMuPDF
import zipfile
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import os

def extract_docx_native(file_bytes):
    """Bulletproof DOCX extraction - handles ALL .docx files"""
    try:
        with zipfile.ZipFile(BytesIO(file_bytes)) as zf:
            # Try document.xml first, fallback to document2.xml or document1.xml
            xml_files = ['word/document.xml', 'word/document2.xml', 'word/document1.xml']
            xml_content = None
            for xml_file in xml_files:
                try:
                    xml_content = zf.read(xml_file).decode('utf-8')
                    break
                except KeyError:
                    continue
            
            if xml_content:
                # Strip XML tags, get clean text
                text = re.sub(r'<[^>]+>', '', xml_content)
                return ' '.join(text.split())
    except:
        pass
    return ""

@st.cache_data
def extract_text(file_bytes, file_name):
    """Universal text extraction"""
    ext = os.path.splitext(file_name)[1].lower()
    text = ""
    
    if ext == '.pdf':
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
    elif ext == '.docx':
        text = extract_docx_native(file_bytes)
    elif ext == '.txt':
        text = file_bytes.decode('utf-8', errors='ignore')
    elif ext in ['.doc', '.rtf']:
        text = file_bytes.decode('utf-8', errors='ignore')
    
    return text

def find_name(text):
    # Extract name from first 300 chars
    name_match = re.search(r'([A-Z][a-z]+ [A-Z][a-z]+)', text[:300])
    if name_match:
        parts = name_match.group(1).split()
        return f"{parts[0].capitalize()} {parts[-1].capitalize()}"
    return "First Last"

def parse_sections(text):
    """Extract sections with keywords"""
    text_lower = text.lower()
    sections = {}
    
    # Summary
    summary = re.search(r'(summary:?)\s*(.*?)(?=(technical skills|education|experience|skills|$))', text_lower+text, re.DOTALL | re.I)
    if summary: sections['summary'] = [line.strip() for line in re.sub(r'[\u2022\u25CF•*-]\s*', '', summary.group(2)).split('\n') if line.strip()]
    
    # Skills
    skills = re.search(r'(technical skills:?)\s*(.*?)(?=(education|experience|$))', text_lower+text, re.DOTALL | re.I)
    if skills: sections['skills'] = [line.strip() for line in re.sub(r'[\u2022\u25CF•*-]\s*', '', skills.group(2)).split('\n') if line.strip()]
    
    # Education/Cert/Training
    edu = re.search(r'(education|certification|training)\s*(.*?)(?=(professional experience|experience|$))', text_lower+text, re.DOTALL | re.I)
    if edu: sections['education'] = [line.strip() for line in re.sub(r'[\u2022\u25CF•*-]\s*', '', edu.group(2)).split('\n') if line.strip()]
    
    # Experience
    exp = re.search(r'(professional experience|experience)\s*(.*)', text_lower+text, re.DOTALL | re.I)
    if exp: sections['experience'] = [line.strip() for line in re.sub(r'[\u2022\u25CF•*-]\s*', '', exp.group(2)).split('\n') if line.strip()]
    
    return sections

def create_docx(name, sections):
    """Create formatted DOCX per exact requirements"""
    doc = Document()
    
    # Set global Times New Roman 10pt
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    
    # 1. Name - Centered, TNR 11 bold
    name_p = doc.add_paragraph(name)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.runs[0]
    run.bold = True
    run.font.size = Pt(11)
    doc.add_paragraph()  # Space
    
    def add_section(title, content):
        # Section heading TNR 10 bold
        p = doc.add_paragraph(title)
        p.runs[0].bold = True
        doc.add_paragraph()  # Space after heading
        
        # Bulleted content
        for line in content[:10]:  # Limit to avoid huge docs
            p = doc.add_paragraph()
            p.add_run('  • ').font.size = Pt(10)
            p.add_run(line).font.size = Pt(10)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
    
    # Add sections per requirements order
    if 'summary' in sections and sections['summary']:
        add_section('Summary', sections['summary'])
    if 'skills' in sections and sections['skills']:
        add_section('Technical Skills', sections['skills'])
    if 'education' in sections and sections['education']:
        add_section('Education, Certification & Training', sections['education'])
    if 'experience' in sections and sections['experience']:
        add_section('Professional Experience', sections['experience'])
    
    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio, f"{name.replace(' ', '')}.docx"

# Streamlit UI
st.title("🔧 Resume Formatter")
st.write("Upload any resume → Get perfectly formatted DOCX")

uploaded_file = st.file_uploader("Choose resume file", type=['pdf','docx','doc','txt','rtf'])

if uploaded_file is not None:
    with st.spinner("Processing..."):
        text = extract_text(uploaded_file.read(), uploaded_file.name)
        name = find_name(text)
        sections = parse_sections(text)
        
        st.success(f"✅ Parsed: **{name}**")
        st.json({k: len(v) for k,v in sections.items() if v})  # Show section counts
        
        doc_io, filename = create_docx(name, sections)
        
        st.download_button(
            label="📥 Download Formatted Resume",
            data=doc_io.getvalue(),
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

st.info("💡 **Deployed on Streamlit Cloud** - Push to GitHub for live URL!")
